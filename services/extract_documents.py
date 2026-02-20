"""Document parsing and extraction utilities for happyRAV."""
from __future__ import annotations

import io
import re
from typing import Dict, Iterable, List, Optional, Tuple

import pdfplumber
from docx import Document as DocxDocument
from PIL import Image

from happyrav.models import (
    DocTag,
    DocumentMeta,
    EducationItem,
    ExperienceItem,
    ExtractedProfile,
    ParseMethod,
    SourceAttribution,
)


MAX_FILE_BYTES = 12 * 1024 * 1024
MAX_SESSION_BYTES = 25 * 1024 * 1024
MAX_SESSION_DOCS = 20

ALLOWED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
}

DOC_TAGS: Tuple[DocTag, ...] = ("cv", "cover_letter", "arbeitszeugnis", "certificate", "other")

_SENSITIVE_BLOCK_RE = re.compile(
    r"-----BEGIN\s[\w\s]+-----[\s\S]*?-----END\s[\w\s]+-----",
    re.MULTILINE,
)


def _sanitize_text(text: str) -> str:
    """Strip PEM-encoded key blocks and other sensitive material from extracted text."""
    return _SENSITIVE_BLOCK_RE.sub("[REDACTED: sensitive key material removed]", text)

MONTH_WORDS = (
    "jan",
    "feb",
    "mar",
    "apr",
    "mai",
    "may",
    "jun",
    "jul",
    "aug",
    "sep",
    "oct",
    "okt",
    "nov",
    "dec",
    "dez",
)


def _extension(filename: str) -> str:
    match = re.search(r"\.([a-zA-Z0-9]+)$", filename or "")
    return f".{match.group(1).lower()}" if match else ""


def is_supported_filename(filename: str) -> bool:
    return _extension(filename) in ALLOWED_EXTENSIONS


def guess_doc_tag(filename: str, provided_tag: Optional[str] = None) -> DocTag:
    if provided_tag in DOC_TAGS:
        return provided_tag  # type: ignore[return-value]

    lower = (filename or "").lower()
    if any(key in lower for key in ("arbeitszeugnis", "zeugnis", "reference", "recommendation")):
        return "arbeitszeugnis"
    if any(key in lower for key in ("cover", "anschreiben", "motivation")):
        return "cover_letter"
    if any(key in lower for key in ("certificate", "cert", "diploma")):
        return "certificate"
    if any(key in lower for key in ("cv", "resume", "lebenslauf")):
        return "cv"
    return "other"


def _has_date_like(line: str) -> bool:
    lower = line.lower()
    return bool(re.search(r"(19|20)\d{2}", lower) or any(month in lower for month in MONTH_WORDS))


def _try_image_ocr(image: Image.Image) -> str:
    try:
        import pytesseract
    except Exception as exc:
        raise RuntimeError("OCR requires pytesseract and tesseract binaries.") from exc
    return pytesseract.image_to_string(image, lang="deu+eng").strip()


def _ocr_pdf_page(content: bytes, page_index: int) -> str:
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError("PDF OCR requires PyMuPDF.") from exc
    with fitz.open(stream=content, filetype="pdf") as doc:
        if page_index < 0 or page_index >= len(doc):
            return ""
        page = doc[page_index]
        matrix = fitz.Matrix(2.2, 2.2)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        image = Image.open(io.BytesIO(pix.tobytes("png")))
        return _try_image_ocr(image)


def extract_text_from_bytes(filename: str, content: bytes) -> Tuple[str, ParseMethod, float]:
    ext = _extension(filename)
    if ext == ".pdf":
        blocks: List[str] = []
        used_ocr = False
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for idx, page in enumerate(pdf.pages):
                text = (page.extract_text() or "").strip()
                if len(text) >= 120:
                    blocks.append(text)
                    continue
                ocr_text = ""
                try:
                    ocr_text = _ocr_pdf_page(content=content, page_index=idx)
                except Exception:
                    ocr_text = ""
                if ocr_text:
                    used_ocr = True
                    blocks.append(ocr_text)
                elif text:
                    blocks.append(text)
        parse_method: ParseMethod = "pdf_text_ocr" if used_ocr else "pdf_text"
        confidence = 0.78 if used_ocr else 0.93
        return _sanitize_text("\n\n".join(blocks).strip()), parse_method, confidence

    if ext == ".docx":
        doc = DocxDocument(io.BytesIO(content))
        blocks = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_cells:
                    blocks.append(" | ".join(row_cells))
        return _sanitize_text("\n".join(blocks).strip()), "docx_text", 0.9

    if ext in {".png", ".jpg", ".jpeg", ".webp"}:
        image = Image.open(io.BytesIO(content))
        text = _try_image_ocr(image)
        return text.strip(), "ocr_image", 0.65

    try:
        decoded = content.decode("utf-8")
    except UnicodeDecodeError:
        decoded = content.decode("latin-1", errors="ignore")
    return _sanitize_text(decoded.strip()), "plain_text", 0.55


def _source_entry(doc_id: str, confidence: float) -> List[SourceAttribution]:
    return [SourceAttribution(doc_id=doc_id, confidence=confidence)]


def _add_source(profile: ExtractedProfile, field_name: str, doc_id: str, confidence: float) -> None:
    profile.source_map.setdefault(field_name, [])
    profile.source_map[field_name].append(SourceAttribution(doc_id=doc_id, confidence=confidence))


def _unique_keep_order(values: Iterable[str]) -> List[str]:
    out: List[str] = []
    seen = set()
    for value in values:
        cleaned = value.strip()
        key = cleaned.lower()
        if cleaned and key not in seen:
            seen.add(key)
            out.append(cleaned)
    return out


def _extract_email(text: str) -> str:
    match = re.search(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", text)
    return match.group(0).strip() if match else ""


def _extract_phone(text: str) -> str:
    match = re.search(r"(\+?\d[\d\s()\/-]{6,}\d)", text)
    return match.group(1).strip() if match else ""


def _extract_linkedin(text: str) -> str:
    match = re.search(r"(https?://(?:www\.)?linkedin\.com/[^\s]+)", text, re.IGNORECASE)
    return match.group(1).strip() if match else ""


def _extract_portfolio(text: str) -> str:
    urls = re.findall(r"(https?://[^\s]+)", text, re.IGNORECASE)
    for url in urls:
        if "linkedin.com" not in url.lower():
            return url.strip()
    for line in text.splitlines():
        clean = line.strip().strip(",.;")
        if not clean or "@" in clean:
            continue
        if re.fullmatch(r"[a-z0-9.-]+\.[a-z]{2,}", clean.lower()) and "linkedin" not in clean.lower():
            return f"https://{clean}"
    return ""


def _extract_name(text: str) -> str:
    for line in text.splitlines()[:20]:
        clean = line.strip()
        if not clean:
            continue
        if re.search(r"(curriculum|lebenslauf|resume|marketing manager|medizinische)", clean, re.IGNORECASE):
            continue
        if re.search(r"@|http|www|tel|phone|kontakt|contact", clean, re.IGNORECASE):
            continue
        words = clean.split()
        if 2 <= len(words) <= 4 and all(re.match(r"^[A-Za-zÄÖÜäöüß'.-]+$", w) for w in words):
            return clean
    return ""


def _extract_skills(text: str) -> List[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    skills: List[str] = []
    for idx, line in enumerate(lines):
        if re.search(r"\b(skills?|kompetenzen|kenntnisse)\b", line, re.IGNORECASE):
            block = lines[idx + 1 : idx + 8]
            joined = ", ".join(block)
            candidates = re.split(r"[,;/•\-\n]+", joined)
            skills.extend([token.strip() for token in candidates if 2 <= len(token.strip()) <= 40])
    if not skills:
        candidates = re.findall(r"\b[A-Za-z][A-Za-z0-9+#.\-]{2,24}\b", text)
        shortlist = [c for c in candidates if c[0].isupper() or "+" in c or "#" in c]
        skills = shortlist[:30]
    return _unique_keep_order(skills)[:25]


def _extract_languages(text: str) -> List[str]:
    lang_tokens = {
        "Deutsch": ["deutsch", "german", "c2", "muttersprache"],
        "English": ["english", "englisch"],
        "Français": ["french", "französisch", "français"],
        "Italiano": ["italian", "italienisch"],
        "Ελληνικά": ["griechisch", "greek"],
    }
    found: List[str] = []
    lower = text.lower()
    for label, keys in lang_tokens.items():
        if any(key in lower for key in keys):
            found.append(label)
    return found


def _extract_achievements(text: str) -> List[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    picks = []
    for line in lines:
        if re.match(r"^[\-•*]\s+", line) and re.search(r"\d|%|kpi|steiger|improv|increase|reduce", line, re.IGNORECASE):
            picks.append(re.sub(r"^[\-•*]\s*", "", line))
    return _unique_keep_order(picks)[:20]


def _line_is_section_heading(line: str) -> bool:
    normalized = re.sub(r"[^a-zäöüß ]+", "", line.lower()).strip()
    tokens = {
        "profil",
        "profile",
        "kontakt",
        "contact",
        "sprachen",
        "sprachkenntnisse",
        "kompetenzen",
        "skills",
        "portfolio",
        "ausbildung",
        "education",
        "berufserfahrung",
        "experience",
        "references",
        "referenzen",
        "zertifikate",
    }
    return normalized in tokens


def _looks_like_experience_header(role: str, company: str) -> bool:
    line = f"{role} {company}".lower()
    bad_tokens = (
        "bsc",
        "msc",
        "bachelor",
        "master",
        "berufsmatur",
        "zertifikat",
        "zertifizierung",
        "hslu",
        "hwz",
        "universität",
        "university",
        "schule",
        "education",
        "ausbildung",
        "experience hslu",
        "teilzeit",
    )
    if any(token in line for token in bad_tokens):
        return False
    role_tokens = (
        "manager",
        "leiter",
        "lead",
        "engineer",
        "entwickler",
        "projekt",
        "consultant",
        "assistant",
        "developer",
        "strategist",
        "analyst",
        "owner",
        "intern",
        "freelance",
    )
    return any(token in role.lower() for token in role_tokens)


def _extract_experience(text: str) -> List[ExperienceItem]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    items: List[ExperienceItem] = []
    inline_pattern = re.compile(
        r"^(?P<role>[^|·]{3,90})\s*(?:\||·)\s*(?P<company>[^|·]{2,90})\s*(?:\||·)\s*(?P<period>.*)$"
    )
    header_pattern = re.compile(r"^(?P<role>[^|·]{3,90})\s*(?:\||·)\s*(?P<company>[^|·]{2,90})$")

    idx = 0
    while idx < len(lines):
        line = lines[idx]
        inline = inline_pattern.match(line)
        if inline and _has_date_like(inline.group("period")) and _looks_like_experience_header(inline.group("role"), inline.group("company")):
            items.append(
                ExperienceItem(
                    role=inline.group("role").strip(),
                    company=inline.group("company").strip(),
                    period=inline.group("period").strip(),
                    achievements=[],
                )
            )
            idx += 1
            continue

        header = header_pattern.match(line)
        if not header:
            idx += 1
            continue
        role = header.group("role").strip()
        company = header.group("company").strip()
        if not _looks_like_experience_header(role, company):
            idx += 1
            continue
        period = ""
        achievements: List[str] = []
        lookahead = idx + 1
        if lookahead < len(lines) and _has_date_like(lines[lookahead]):
            period = lines[lookahead]
            lookahead += 1
        while lookahead < len(lines):
            candidate = lines[lookahead]
            if _line_is_section_heading(candidate):
                break
            if header_pattern.match(candidate):
                break
            if len(candidate) < 3:
                lookahead += 1
                continue
            if candidate.startswith(("•", "-", "*")) or re.search(r"\d|%|result|ergebnis|lead|optim|steiger", candidate, re.IGNORECASE):
                cleaned = re.sub(r"^[•\-*]\s*", "", candidate)
                achievements.append(cleaned)
            if len(achievements) >= 5:
                break
            lookahead += 1
        items.append(
            ExperienceItem(
                role=role,
                company=company,
                period=period,
                achievements=_unique_keep_order(achievements)[:5],
            )
        )
        idx = max(lookahead, idx + 1)

    if items:
        deduped: List[ExperienceItem] = []
        seen = set()
        for item in items:
            key = (item.role.lower(), item.company.lower(), item.period.lower())
            if key in seen:
                continue
            seen.add(key)
            deduped.append(item)
        return deduped[:12]

    for line in lines:
        if re.search(r"(19|20)\d{2}", line) and re.search(
            r"\b(manager|lead|leiter|engineer|projekt|consultant|assistant|specialist|developer)\b",
            line,
            re.IGNORECASE,
        ):
            items.append(ExperienceItem(role=line, company="", period="", achievements=[]))
        elif re.search(
            r"\b(manager|lead|leiter|engineer|projekt|consultant|assistant|specialist|developer|owner|strategist)\b",
            line,
            re.IGNORECASE,
        ) and not _line_is_section_heading(line):
            items.append(ExperienceItem(role=line, company="", period="", achievements=[]))
    return items[:10]


def _extract_education(text: str) -> List[EducationItem]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    items: List[EducationItem] = []
    degree_pattern = re.compile(r"\b(BSc|MSc|Bachelor|Master|CAS|Diplom|PhD|Doktor|Berufsmaturität)\b", re.IGNORECASE)
    for idx, line in enumerate(lines):
        if not degree_pattern.search(line):
            continue
        school = ""
        period = ""
        if idx + 1 < len(lines) and not degree_pattern.search(lines[idx + 1]) and not _line_is_section_heading(lines[idx + 1]):
            if _has_date_like(lines[idx + 1]):
                period = lines[idx + 1]
            else:
                school = lines[idx + 1]
        if idx + 2 < len(lines) and not period and _has_date_like(lines[idx + 2]):
            period = lines[idx + 2]
        items.append(EducationItem(degree=line, school=school, period=period))
    return items[:8]


def extract_profile_fragment(
    text: str,
    doc_id: str,
    confidence: float,
) -> ExtractedProfile:
    fragment = ExtractedProfile()
    if not text:
        return fragment

    email = _extract_email(text)
    phone = _extract_phone(text)
    linkedin = _extract_linkedin(text)
    portfolio = _extract_portfolio(text)
    full_name = _extract_name(text)
    skills = _extract_skills(text)
    languages = _extract_languages(text)
    achievements = _extract_achievements(text)
    experience = _extract_experience(text)
    education = _extract_education(text)

    summary_lines = [line.strip() for line in text.splitlines() if line.strip()][:2]
    summary = " ".join(summary_lines)[:400]

    if full_name:
        fragment.full_name = full_name
        _add_source(fragment, "full_name", doc_id, confidence)
    if email:
        fragment.email = email
        _add_source(fragment, "email", doc_id, confidence)
    if phone:
        fragment.phone = phone
        _add_source(fragment, "phone", doc_id, confidence)
    if linkedin:
        fragment.linkedin = linkedin
        _add_source(fragment, "linkedin", doc_id, confidence)
    if portfolio:
        fragment.portfolio = portfolio
        _add_source(fragment, "portfolio", doc_id, confidence)
    if summary:
        fragment.summary = summary
        _add_source(fragment, "summary", doc_id, confidence)

    fragment.skills = skills
    fragment.languages = languages
    fragment.achievements = achievements
    fragment.experience = experience
    fragment.education = education

    if skills:
        _add_source(fragment, "skills", doc_id, confidence)
    if experience:
        _add_source(fragment, "experience", doc_id, confidence)
    if education:
        _add_source(fragment, "education", doc_id, confidence)

    return fragment


def merge_profiles(base: ExtractedProfile, fragment: ExtractedProfile) -> ExtractedProfile:
    merged = base.model_copy(deep=True)
    scalar_fields = [
        "full_name",
        "headline",
        "email",
        "phone",
        "location",
        "linkedin",
        "portfolio",
        "summary",
    ]
    for field in scalar_fields:
        current = getattr(merged, field)
        incoming = getattr(fragment, field)
        if not current and incoming:
            setattr(merged, field, incoming)

    merged.skills = _unique_keep_order([*merged.skills, *fragment.skills])[:30]
    merged.languages = _unique_keep_order([*merged.languages, *fragment.languages])[:10]
    merged.achievements = _unique_keep_order([*merged.achievements, *fragment.achievements])[:30]

    exp_seen = {(item.role.lower(), item.company.lower(), item.period.lower()) for item in merged.experience}
    for item in fragment.experience:
        key = (item.role.lower(), item.company.lower(), item.period.lower())
        if key not in exp_seen and item.role:
            exp_seen.add(key)
            merged.experience.append(item)

    edu_seen = {(item.degree.lower(), item.school.lower(), item.period.lower()) for item in merged.education}
    for item in fragment.education:
        key = (item.degree.lower(), item.school.lower(), item.period.lower())
        if key not in edu_seen and item.degree:
            edu_seen.add(key)
            merged.education.append(item)

    for field_name, sources in fragment.source_map.items():
        merged.source_map.setdefault(field_name, [])
        merged.source_map[field_name].extend(sources)

    return merged


def build_document_meta(
    doc_id: str,
    filename: str,
    mime: str,
    tag: DocTag,
    parse_method: ParseMethod,
    confidence: float,
    size_bytes: int,
    text: str,
) -> DocumentMeta:
    excerpt = (text[:240] + "...") if len(text) > 240 else text
    return DocumentMeta(
        doc_id=doc_id,
        filename=filename,
        mime=mime,
        tag=tag,
        parse_method=parse_method,
        confidence=confidence,
        size_bytes=size_bytes,
        text_excerpt=excerpt,
    )
